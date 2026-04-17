#!/usr/bin/env python3
"""
Word Document Extractor
Extract text content from Word (.docx) documents
"""

import sys
import argparse
from docx import Document

def extract_text(docx_path, output_path=None, preserve_structure=True):
    """
    Extract text from Word document

    Args:
        docx_path: Path to the .docx file
        output_path: Optional output file path
        preserve_structure: Whether to preserve document structure
    """
    try:
        doc = Document(docx_path)

        if preserve_structure:
            # Extract with structure (paragraphs, headings, lists)
            content = []

            for para in doc.paragraphs:
                # Skip empty paragraphs
                if not para.text.strip():
                    continue

                # Detect heading style
                style_name = para.style.name if para.style else "Normal"

                # Format based on style
                if "Heading 1" in style_name:
                    content.append(f"\n# {para.text}\n")
                elif "Heading 2" in style_name:
                    content.append(f"\n## {para.text}\n")
                elif "Heading 3" in style_name:
                    content.append(f"\n### {para.text}\n")
                elif "Heading 4" in style_name:
                    content.append(f"\n#### {para.text}\n")
                else:
                    content.append(f"{para.text}\n")

            # Extract tables
            if doc.tables:
                content.append("\n## Tables\n\n")

            for table_idx, table in enumerate(doc.tables):
                content.append(f"### Table {table_idx + 1}\n")

                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_idx == 0:
                        content.append(f"| {row_text} |\n")
                        content.append(f"| {'--- | ' * len(row.cells)}\n")
                    else:
                        content.append(f"| {row_text} |\n")
                content.append("\n")

            # Join content
            text = "".join(content)

        else:
            # Extract plain text only
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # Output
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"✓ Content extracted to: {output_path}", file=sys.stderr)
        else:
            print(text)

        return text

    except FileNotFoundError:
        print(f"Error: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error extracting Word document: {e}", file=sys.stderr)
        sys.exit(1)

def to_markdown(docx_path, output_path):
    """
    Convert Word document to Markdown format

    Args:
        docx_path: Path to the .docx file
        output_path: Path for the output .md file
    """
    extract_text(docx_path, output_path, preserve_structure=True)

def main():
    parser = argparse.ArgumentParser(description='Extract text from Word documents')
    parser.add_argument('docx_file', help='Path to the Word (.docx) file')
    parser.add_argument('-o', '--output', help='Output file path (optional)')
    parser.add_argument('--plain', action='store_true', help='Extract plain text without structure')
    parser.add_argument('--md', '--markdown', action='store_true', help='Convert to Markdown format')

    args = parser.parse_args()

    if args.md or not args.plain:
        # Default to Markdown format
        output_path = args.output if args.output else None
        extract_text(args.docx_file, output_path, preserve_structure=True)
    else:
        extract_text(args.docx_file, args.output, preserve_structure=False)

if __name__ == "__main__":
    main()
